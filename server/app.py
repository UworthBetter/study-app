import os
import io
import re
import requests
from functools import wraps
from flask import Flask, request, jsonify, g
from flask_cors import CORS
from pdf_parser import parse_pdf
from database import init_db
from models import (
    get_or_create_user, update_user_profile, create_session, get_session,
    delete_session, get_all_official_banks, get_official_bank_by_id,
    add_user_favorite, is_already_favorited,
)

app = Flask(__name__)
CORS(app)

WX_APPID = os.environ.get('WX_APPID', '')
WX_SECRET = os.environ.get('WX_SECRET', '')

# 启动时初始化数据库
init_db()


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.headers.get('Authorization', '')
        if not auth.startswith('Bearer '):
            return jsonify({"error": "未登录"}), 401
        token = auth[7:]
        session = get_session(token)
        if not session:
            return jsonify({"error": "登录已过期"}), 401
        user = get_or_create_user(session["openid"])
        g.current_user = user
        g.token = token
        return f(*args, **kwargs)
    return decorated


def parse_docx(docx_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs.append(' '.join(row_texts))
        result = '\n'.join(paragraphs)
        result = re.sub(r'\n{3,}', '\n\n', result)
        result = re.sub(r' +', ' ', result)
        return result.strip()
    except ImportError:
        raise Exception("服务器未安装 python-docx，请运行: pip install python-docx")


# ========== 认证 API ==========

@app.route("/api/auth/wx-login", methods=["POST"])
def wx_login():
    data = request.get_json()
    if not data or 'code' not in data:
        return jsonify({"error": "缺少 code 参数"}), 400

    code = data['code']

    # 开发模式：未配置 WX_APPID/WX_SECRET 时，用 code 的 hash 生成伪 openid
    if not WX_APPID or not WX_SECRET:
        import hashlib
        mock_openid = 'dev_' + hashlib.md5(code.encode()).hexdigest()[:16]
        print(f"[开发模式] 未配置 WX_APPID/WX_SECRET，使用模拟 openid: {mock_openid}")
        user = get_or_create_user(mock_openid)
        is_new = user['created_at'] == user['last_login']
        token = create_session(mock_openid)
        return jsonify({
            "token": token,
            "user": {
                "openid": mock_openid,
                "nickname": user['nickname'],
                "avatarUrl": user['avatar_url'],
            },
            "isNewUser": is_new,
        })

    url = (
        f"https://api.weixin.qq.com/sns/jscode2session"
        f"?appid={WX_APPID}&secret={WX_SECRET}"
        f"&js_code={code}&grant_type=authorization_code"
    )
    try:
        resp = requests.get(url, timeout=10)
        wx_data = resp.json()
    except Exception as e:
        return jsonify({"error": f"微信接口调用失败: {str(e)}"}), 500

    if 'errcode' in wx_data and wx_data['errcode'] != 0:
        return jsonify({"error": f"微信登录失败: {wx_data.get('errmsg', '未知错误')}"}), 400

    openid = wx_data.get('openid')
    if not openid:
        return jsonify({"error": "未获取到 openid"}), 400

    user = get_or_create_user(openid)
    is_new = user['created_at'] == user['last_login']
    token = create_session(openid)

    return jsonify({
        "token": token,
        "user": {
            "openid": openid,
            "nickname": user['nickname'],
            "avatarUrl": user['avatar_url'],
        },
        "isNewUser": is_new,
    })


@app.route("/api/auth/update-profile", methods=["POST"])
@login_required
def update_profile():
    data = request.get_json()
    if not data:
        return jsonify({"error": "缺少参数"}), 400
    nickname = data.get('nickname', '')
    avatar_url = data.get('avatarUrl', '')
    update_user_profile(g.current_user['openid'], nickname, avatar_url)
    return jsonify({"success": True})


@app.route("/api/auth/logout", methods=["POST"])
@login_required
def logout():
    delete_session(g.token)
    return jsonify({"success": True})


# ========== 官方题库 API ==========

@app.route("/api/official-banks", methods=["GET"])
def list_official_banks():
    banks = get_all_official_banks()
    return jsonify({"banks": banks})


@app.route("/api/official-banks/<int:bank_id>", methods=["GET"])
def get_official_bank(bank_id):
    bank = get_official_bank_by_id(bank_id)
    if not bank:
        return jsonify({"error": "题库不存在"}), 404
    return jsonify({"bank": bank})


@app.route("/api/official-banks/<int:bank_id>/favorite", methods=["POST"])
@login_required
def favorite_bank(bank_id):
    bank = get_official_bank_by_id(bank_id)
    if not bank:
        return jsonify({"error": "题库不存在"}), 404

    if is_already_favorited(g.current_user['id'], bank_id):
        return jsonify({"error": "已收藏过该题库"}), 400

    if add_user_favorite(g.current_user['id'], bank_id):
        return jsonify({"success": True, "bank": bank})
    return jsonify({"error": "收藏失败"}), 500


# ========== 文档解析 API ==========

@app.route("/api/parse-pdf", methods=["POST"])
def parse_pdf_endpoint():
    if "file" not in request.files:
        return jsonify({"error": "没有上传文件"}), 400
    file = request.files["file"]
    filename = file.filename.lower()
    try:
        file_bytes = file.read()
        if filename.endswith(".pdf"):
            text = parse_pdf(file_bytes)
        elif filename.endswith(".docx") or filename.endswith(".doc"):
            text = parse_docx(file_bytes)
        else:
            return jsonify({"error": "只支持 PDF 和 Word 文件"}), 400
        if not text or not text.strip():
            return jsonify({"error": "未能提取到文字内容"}), 400
        print(f"提取文本长度: {len(text)}")
        return jsonify({"text": text, "length": len(text)})
    except Exception as e:
        print(f"解析错误: {str(e)}")
        return jsonify({"error": f"解析失败: {str(e)}"}), 500


@app.route("/api/parse-document", methods=["POST"])
def parse_document_endpoint():
    return parse_pdf_endpoint()


@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


if __name__ == "__main__":
    print("服务器启动在 http://0.0.0.0:5000")
    app.run(host="0.0.0.0", port=5000, debug=True)
